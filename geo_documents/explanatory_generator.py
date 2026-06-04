from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.text import WD_COLOR_INDEX

from geo_documents.llm_ollama import DEFAULT_MODEL, DEFAULT_OLLAMA_HOST, generate_text
from geo_documents.template_model import ExplanationTemplate, TemplateBlock


@dataclass
class GeneratedBlock:
    source: TemplateBlock
    text: str


@dataclass
class GenerationResult:
    template: ExplanationTemplate
    blocks: list[GeneratedBlock]

    @property
    def text(self) -> str:
        return "\n\n".join(block.text for block in self.blocks if block.text.strip())


def build_block_task(template: ExplanationTemplate, block: TemplateBlock) -> str:
    note = block.note.strip() or "Найти подходящую информацию в контексте документа."
    return (
        "Заполни один фрагмент пояснительной записки по шаблону.\n\n"
        f"Название шаблона: {template.name}\n"
        f"Текст-заготовка фрагмента: {block.text}\n"
        f"Инструкция/заметка к фрагменту: {note}\n\n"
        "Требования:\n"
        "- верни только готовый текст фрагмента;\n"
        "- не добавляй комментарии и заголовки, если они не требуются фрагментом;\n"
        "- используй только факты из контекста документа;\n"
        "- если данных нет, напиши [уточнить];\n"
        "- стиль: официальный технический русский язык."
    )


def generate_explanatory_note(
    template: ExplanationTemplate,
    context_paths: list[str | Path],
    *,
    model: str = DEFAULT_MODEL,
    host: str = DEFAULT_OLLAMA_HOST,
    timeout: int = 300,
    low_memory: bool = True,
    max_chars_per_file: int = 8_000,
    max_total_chars: int = 12_000,
) -> GenerationResult:
    blocks: list[GeneratedBlock] = []

    for block in template.blocks:
        if block.type == "plain":
            # Неразмеченный текст служит только каркасом шаблона в редакторе.
            # В итоговую пояснительную записку он не попадает.
            continue
        if block.type == "fixed":
            blocks.append(GeneratedBlock(source=block, text=block.text))
            continue

        task = build_block_task(template, block)
        generated = generate_text(
            task=task,
            context_paths=context_paths,
            model=model,
            host=host,
            timeout=timeout,
            low_memory=low_memory,
            max_chars_per_file=max_chars_per_file,
            max_total_chars=max_total_chars,
        )
        blocks.append(GeneratedBlock(source=block, text=generated))

    return GenerationResult(template=template, blocks=blocks)


def _add_text_block(doc: Document, block: GeneratedBlock, *, keep_highlight: bool) -> None:
    text = block.text.strip()
    if not text:
        return

    for idx, paragraph_text in enumerate(text.splitlines()):
        if not paragraph_text.strip():
            continue
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(paragraph_text.strip())
        if keep_highlight and block.source.type == "fixed":
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
        elif keep_highlight and block.source.type == "generated":
            run.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
        if idx == 0 and block.source.note:
            paragraph.style = doc.styles["Normal"]


def save_generation_result_docx(
    result: GenerationResult,
    output_path: str | Path,
    *,
    keep_highlight: bool = False,
) -> Path:
    doc = Document()
    for block in result.blocks:
        _add_text_block(doc, block, keep_highlight=keep_highlight)
    output = Path(output_path)
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output))
    return output
