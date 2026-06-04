from __future__ import annotations

import json
import urllib.error
import urllib.request
from dataclasses import dataclass
from pathlib import Path
from typing import Any

import fitz  # PyMuPDF
from docx import Document


DEFAULT_OLLAMA_HOST = "http://127.0.0.1:11434"
# Gemma4 — мультимодальная линейка, даже e2b часто не влезает в 8 ГБ RAM.
# Для слабых ПК: gemma3:1b, gemma2:2b или phi3:mini (только текст).
DEFAULT_MODEL = "gemma3:1b"
LOW_RAM_MODEL_HINT = "gemma3:1b, gemma2:2b или phi3:mini"
DEFAULT_SYSTEM_PROMPT = (
    "Ты локальный помощник для подготовки инженерно-геологических отчетов. "
    "Пиши по-русски, деловым техническим стилем. Используй только предоставленный "
    "контекст документа и явно не выдумывай факты, которых в контексте нет."
)
SUPPORTED_CONTEXT_EXTS = {".docx", ".pdf", ".txt", ".md"}


class OllamaError(RuntimeError):
    """Ошибка локального LLM/Ollama слоя."""


@dataclass(frozen=True)
class DocumentContext:
    path: Path
    text: str
    truncated: bool = False


def _normalize_host(host: str) -> str:
    return host.rstrip("/")


def _truncate_text(text: str, max_chars: int) -> tuple[str, bool]:
    text = "\n".join(line.strip() for line in text.splitlines() if line.strip())
    if len(text) <= max_chars:
        return text, False
    return text[:max_chars].rstrip() + "\n[...контекст обрезан...]", True


def _read_docx_text(path: Path) -> str:
    doc = Document(str(path))
    parts: list[str] = []

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            row_text = " | ".join(cell for cell in cells if cell)
            if row_text:
                parts.append(row_text)

    return "\n".join(parts)


def _read_pdf_text(path: Path, *, max_pages: int | None = None) -> str:
    parts: list[str] = []
    pdf = fitz.open(path)
    try:
        page_count = len(pdf) if max_pages is None else min(len(pdf), max_pages)
        for page_idx in range(page_count):
            text = pdf[page_idx].get_text("text").strip()
            if text:
                parts.append(f"[Страница {page_idx + 1}]\n{text}")
    finally:
        pdf.close()
    return "\n".join(parts)


def _read_plain_text(path: Path) -> str:
    for encoding in ("utf-8", "cp1251"):
        try:
            return path.read_text(encoding=encoding)
        except UnicodeDecodeError:
            continue
    return path.read_text(encoding="utf-8", errors="replace")


def read_document_context(
    paths: list[str | Path],
    *,
    max_chars_per_file: int = 12_000,
    max_pages_per_pdf: int | None = 20,
) -> list[DocumentContext]:
    """Извлекает текстовый контекст из DOCX/PDF/TXT/MD для локальной LLM."""
    contexts: list[DocumentContext] = []

    for raw_path in paths:
        path = Path(raw_path)
        ext = path.suffix.lower()
        if not path.is_file() or ext not in SUPPORTED_CONTEXT_EXTS:
            continue

        if ext == ".docx":
            text = _read_docx_text(path)
        elif ext == ".pdf":
            text = _read_pdf_text(path, max_pages=max_pages_per_pdf)
        else:
            text = _read_plain_text(path)

        text, truncated = _truncate_text(text, max_chars_per_file)
        if text:
            contexts.append(DocumentContext(path=path, text=text, truncated=truncated))

    return contexts


def build_context_prompt(
    task: str,
    contexts: list[DocumentContext],
    *,
    max_total_chars: int = 30_000,
) -> str:
    """Собирает пользовательский промпт с контекстом документов."""
    blocks: list[str] = []
    used_chars = 0

    for ctx in contexts:
        title = f"### Файл: {ctx.path.name}"
        text = ctx.text
        remaining = max_total_chars - used_chars - len(title) - 16
        if remaining <= 0:
            break
        if len(text) > remaining:
            text = text[:remaining].rstrip() + "\n[...общий контекст обрезан...]"
        blocks.append(f"{title}\n{text}")
        used_chars += len(title) + len(text)

    context_text = "\n\n".join(blocks) if blocks else "Контекст документа не передан или не извлечен."
    return (
        "Задача пользователя:\n"
        f"{task.strip()}\n\n"
        "Контекст документов:\n"
        f"{context_text}\n\n"
        "Сформируй готовый текст. Не добавляй пояснения о том, что ты ИИ."
    )


def _format_ollama_http_error(host: str, error: urllib.error.HTTPError) -> str:
    body = ""
    try:
        body = error.read().decode("utf-8", errors="replace")
        detail = json.loads(body).get("error", body) if body else ""
    except (json.JSONDecodeError, AttributeError):
        detail = body
    detail_lower = str(detail).lower()
    if "allocate" in detail_lower or "out of memory" in detail_lower or "unable to allocate" in detail_lower:
        return (
            f"Недостаточно оперативной памяти для модели ({error.code}). "
            f"Gemma4 (e2b/e4b) тяжёлая даже для текста. Закройте лишние программы, "
            f"затем: ollama pull {LOW_RAM_MODEL_HINT.split(',')[0]} и "
            f"--model {LOW_RAM_MODEL_HINT.split(',')[0]} --low-memory. "
            f"Детали: {detail}"
        )
    return f"Ollama вернула ошибку HTTP {error.code}: {detail or error.reason}"


def _post_json(host: str, endpoint: str, payload: dict[str, Any], *, timeout: int) -> dict[str, Any]:
    url = f"{_normalize_host(host)}{endpoint}"
    data = json.dumps(payload, ensure_ascii=False).encode("utf-8")
    request = urllib.request.Request(
        url,
        data=data,
        headers={"Content-Type": "application/json"},
        method="POST",
    )
    try:
        with urllib.request.urlopen(request, timeout=timeout) as response:
            return json.loads(response.read().decode("utf-8"))
    except urllib.error.HTTPError as e:
        raise OllamaError(_format_ollama_http_error(host, e)) from e
    except urllib.error.URLError as e:
        if isinstance(e.reason, ConnectionRefusedError) or "refused" in str(e).lower():
            raise OllamaError(
                "Ollama недоступна. Запустите приложение Ollama или проверьте адрес "
                f"{_normalize_host(host)}."
            ) from e
        raise OllamaError(f"Ошибка соединения с Ollama: {e}") from e
    except json.JSONDecodeError as e:
        raise OllamaError("Ollama вернула некорректный JSON-ответ.") from e


def list_local_models(host: str = DEFAULT_OLLAMA_HOST, *, timeout: int = 10) -> list[str]:
    url = f"{_normalize_host(host)}/api/tags"
    try:
        with urllib.request.urlopen(url, timeout=timeout) as response:
            payload = json.loads(response.read().decode("utf-8"))
    except urllib.error.URLError as e:
        raise OllamaError(
            "Ollama недоступна. Запустите Ollama и проверьте адрес "
            f"{_normalize_host(host)}."
        ) from e
    except json.JSONDecodeError as e:
        raise OllamaError("Ollama вернула некорректный JSON-ответ.") from e

    return [item.get("name", "") for item in payload.get("models", []) if item.get("name")]


def generate_text(
    *,
    task: str,
    context_paths: list[str | Path],
    model: str = DEFAULT_MODEL,
    host: str = DEFAULT_OLLAMA_HOST,
    system_prompt: str = DEFAULT_SYSTEM_PROMPT,
    temperature: float = 0.2,
    timeout: int = 180,
    max_chars_per_file: int = 12_000,
    max_total_chars: int = 30_000,
    num_ctx: int | None = None,
    low_memory: bool = False,
) -> str:
    """Генерирует текст через локальную Ollama-модель с учетом контекста документов."""
    if low_memory:
        max_chars_per_file = min(max_chars_per_file, 4_000)
        max_total_chars = min(max_total_chars, 8_000)
        if num_ctx is None:
            num_ctx = 2048
    elif num_ctx is None:
        num_ctx = 8192

    contexts = read_document_context(context_paths, max_chars_per_file=max_chars_per_file)
    prompt = build_context_prompt(task, contexts, max_total_chars=max_total_chars)
    payload = {
        "model": model,
        "stream": False,
        "messages": [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": prompt},
        ],
        "options": {"temperature": temperature, "num_ctx": num_ctx},
    }
    response = _post_json(host, "/api/chat", payload, timeout=timeout)

    message = response.get("message")
    if not isinstance(message, dict) or not message.get("content"):
        raise OllamaError("Ollama не вернула текст ответа.")
    return str(message["content"]).strip()
