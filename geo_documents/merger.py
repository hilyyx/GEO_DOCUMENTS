from __future__ import annotations

import io
import shutil
import tempfile
import traceback
import zipfile
from xml.etree import ElementTree
from pathlib import Path

import fitz  # PyMuPDF
from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.image.image import Image as DocxImage
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.section import Section
from docx.shared import Inches
from docxcompose.composer import Composer

from geo_documents.docx_fit import document_ends_with_page_forcing_break, remove_empty_pages, section_content_inches
from geo_documents.libreoffice import docx_to_pdf, find_soffice

RASTER_IMAGE_EXTS = {".png", ".jpg", ".jpeg"}
VECTOR_IMAGE_EXTS = {".svg", ".dvg"}
IMAGE_EXTS = RASTER_IMAGE_EXTS | VECTOR_IMAGE_EXTS
_XML_SPACE = "{http://www.w3.org/XML/1998/namespace}space"
_APP_PROPS_NS = {"ep": "http://schemas.openxmlformats.org/officeDocument/2006/extended-properties"}


def _set_update_fields_on_open(doc: Document) -> None:
    settings = doc.settings.element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def _field_char(kind: str):
    run = OxmlElement("w:r")
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), kind)
    run.append(fld_char)
    return run


def _field_instruction(text: str):
    run = OxmlElement("w:r")
    instr = OxmlElement("w:instrText")
    instr.set(_XML_SPACE, "preserve")
    instr.text = text
    run.append(instr)
    return run


def _field_cached_text(text: str):
    run = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    return run


def _clear_paragraph(paragraph) -> None:
    for child in list(paragraph._p):
        paragraph._p.remove(child)


def _numbering_paragraph(container):
    paragraph = container.paragraphs[0] if container.paragraphs and not container.paragraphs[0].text else container.add_paragraph()
    _clear_paragraph(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    return paragraph


def _add_page_field(paragraph, *, fallback: str = "1") -> None:
    paragraph._p.append(_field_char("begin"))
    paragraph._p.append(_field_instruction(" PAGE "))
    paragraph._p.append(_field_char("separate"))
    paragraph._p.append(_field_cached_text(fallback))
    paragraph._p.append(_field_char("end"))


def _add_global_page_field(paragraph, *, offset: int) -> None:
    if offset <= 0:
        _add_page_field(paragraph)
        return

    paragraph._p.append(_field_char("begin"))
    paragraph._p.append(_field_instruction(" = "))
    paragraph._p.append(_field_char("begin"))
    paragraph._p.append(_field_instruction(" PAGE "))
    paragraph._p.append(_field_char("separate"))
    paragraph._p.append(_field_cached_text("1"))
    paragraph._p.append(_field_char("end"))
    paragraph._p.append(_field_instruction(f" + {offset} "))
    paragraph._p.append(_field_char("separate"))
    paragraph._p.append(_field_cached_text(str(offset + 1)))
    paragraph._p.append(_field_char("end"))


def _set_section_page_restart(section: Section, *, start: int | None) -> None:
    sect_pr = section._sectPr
    for existing in list(sect_pr.findall(qn("w:pgNumType"))):
        sect_pr.remove(existing)
    if start is None:
        return
    pg_num_type = OxmlElement("w:pgNumType")
    pg_num_type.set(qn("w:start"), str(start))
    sect_pr.append(pg_num_type)


def _apply_section_numbering(section: Section, *, global_offset: int, restart_local: bool) -> None:
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    _set_section_page_restart(section, start=1 if restart_local else None)

    header_paragraph = _numbering_paragraph(section.header)
    _add_global_page_field(header_paragraph, offset=global_offset)

    footer_paragraph = _numbering_paragraph(section.footer)
    _add_page_field(footer_paragraph)


def _apply_page_numbering(doc: Document, part_sections: list[tuple[int, int, int]]) -> None:
    if not doc.sections:
        return
    _set_update_fields_on_open(doc)
    for start_idx, end_idx, global_offset in part_sections:
        for idx in range(start_idx, min(end_idx, len(doc.sections) - 1) + 1):
            _apply_section_numbering(
                doc.sections[idx],
                global_offset=global_offset,
                restart_local=idx == start_idx,
            )


def _docx_extended_page_count(path: Path) -> int | None:
    try:
        with zipfile.ZipFile(path) as zf:
            with zf.open("docProps/app.xml") as app_xml:
                root = ElementTree.parse(app_xml).getroot()
    except (OSError, KeyError, ElementTree.ParseError, zipfile.BadZipFile):
        return None

    pages_el = root.find("ep:Pages", _APP_PROPS_NS)
    if pages_el is None or not pages_el.text:
        return None
    try:
        pages = int(pages_el.text)
    except ValueError:
        return None
    return pages if pages > 0 else None


def _pdf_page_count(path: Path) -> int:
    src = fitz.open(path)
    try:
        return max(1, len(src))
    finally:
        src.close()


def _docx_page_count(path: Path, *, soffice: str | None, warnings: list[str]) -> int:
    if soffice:
        tmp_dir = Path(tempfile.mkdtemp(prefix="geo_pages_"))
        try:
            pdf_path = docx_to_pdf(soffice, path, tmp_dir)
            return _pdf_page_count(pdf_path)
        except Exception as e:
            warnings.append(f"Не удалось точно посчитать страницы DOCX ({path.name}): {e}")
        finally:
            shutil.rmtree(tmp_dir, ignore_errors=True)

    pages = _docx_extended_page_count(path)
    if pages is not None:
        return pages
    warnings.append(f"Нумерация для DOCX может быть неточной, не удалось посчитать страницы: {path.name}")
    return 1


def _source_page_count(path: Path, *, soffice: str | None, warnings: list[str]) -> int:
    ext = path.suffix.lower()
    if ext == ".pdf":
        try:
            return _pdf_page_count(path)
        except Exception as e:
            warnings.append(f"Не удалось посчитать страницы PDF ({path.name}): {e}")
            return 1
    if ext == ".docx":
        return _docx_page_count(path, soffice=soffice, warnings=warnings)
    if ext in IMAGE_EXTS:
        return 1
    return 1


def _prepend_heading(doc: Document, text: str) -> None:
    h = doc.add_heading(text, level=2)
    el = h._element
    body = doc.element.body
    body.remove(el)
    body.insert(0, el)


def _append_page_break(doc: Document) -> None:
    if document_ends_with_page_forcing_break(doc):
        return
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def _has_body_content(doc: Document) -> bool:
    return any(child.tag.rsplit("}", 1)[-1] != "sectPr" for child in doc.element.body)


def _section_is_landscape(doc: Document) -> bool:
    section = doc.sections[-1]
    page_width, page_height = _section_page_size(section, fallback=None)
    return page_width > page_height


def _safe_length(value, fallback):
    return value if value is not None else fallback


def _section_page_size(section: Section, *, fallback: Section | None) -> tuple:
    fallback_width = fallback.page_width if fallback is not None and fallback.page_width is not None else Inches(8.27)
    fallback_height = fallback.page_height if fallback is not None and fallback.page_height is not None else Inches(11.69)
    return (
        _safe_length(section.page_width, fallback_width),
        _safe_length(section.page_height, fallback_height),
    )


def _copy_section_geometry(target: Section, source: Section) -> None:
    page_width, page_height = _section_page_size(source, fallback=target)
    target.orientation = WD_ORIENT.LANDSCAPE if page_width > page_height else WD_ORIENT.PORTRAIT
    target.page_width = page_width
    target.page_height = page_height
    target.left_margin = _safe_length(source.left_margin, target.left_margin or Inches(1))
    target.right_margin = _safe_length(source.right_margin, target.right_margin or Inches(1))
    target.top_margin = _safe_length(source.top_margin, target.top_margin or Inches(1))
    target.bottom_margin = _safe_length(source.bottom_margin, target.bottom_margin or Inches(1))
    target.header_distance = _safe_length(source.header_distance, target.header_distance or Inches(0.5))
    target.footer_distance = _safe_length(source.footer_distance, target.footer_distance or Inches(0.5))
    target.gutter = _safe_length(source.gutter, target.gutter or 0)


def _section_geometry_matches(target: Section, source: Section) -> bool:
    target_width, target_height = _section_page_size(target, fallback=None)
    source_width, source_height = _section_page_size(source, fallback=target)
    return (
        target_width == source_width
        and target_height == source_height
        and target.left_margin == _safe_length(source.left_margin, target.left_margin)
        and target.right_margin == _safe_length(source.right_margin, target.right_margin)
        and target.top_margin == _safe_length(source.top_margin, target.top_margin)
        and target.bottom_margin == _safe_length(source.bottom_margin, target.bottom_margin)
    )


def _ensure_section_like(doc: Document, source: Document) -> None:
    if not source.sections:
        return
    source_section = source.sections[0]
    if not _has_body_content(doc):
        _copy_section_geometry(doc.sections[-1], source_section)
        return
    if _section_geometry_matches(doc.sections[-1], source_section):
        return
    doc.add_section(WD_SECTION.NEW_PAGE)
    _copy_section_geometry(doc.sections[-1], source_section)


def _start_new_part_section_like(doc: Document, source: Document, *, page_break: bool) -> None:
    if not source.sections:
        if _has_body_content(doc):
            doc.add_section(WD_SECTION.NEW_PAGE if page_break else WD_SECTION.CONTINUOUS)
        return
    source_section = source.sections[0]
    if not _has_body_content(doc):
        _copy_section_geometry(doc.sections[-1], source_section)
        return
    doc.add_section(WD_SECTION.NEW_PAGE if page_break else WD_SECTION.CONTINUOUS)
    _copy_section_geometry(doc.sections[-1], source_section)


def _set_section_orientation(doc: Document, *, landscape: bool) -> None:
    section = doc.sections[-1]
    page_width, page_height = _section_page_size(section, fallback=None)
    already_landscape = page_width > page_height
    if already_landscape == landscape:
        return
    section.orientation = WD_ORIENT.LANDSCAPE if landscape else WD_ORIENT.PORTRAIT
    section.page_width, section.page_height = page_height, page_width


def _ensure_page_orientation(doc: Document, *, landscape: bool) -> None:
    if not _has_body_content(doc):
        _set_section_orientation(doc, landscape=landscape)
        return
    if _section_is_landscape(doc) == landscape:
        return
    doc.add_section(WD_SECTION.NEW_PAGE)
    _set_section_orientation(doc, landscape=landscape)


def _start_new_part_section_for_orientation(doc: Document, *, landscape: bool, page_break: bool) -> None:
    if _has_body_content(doc):
        doc.add_section(WD_SECTION.NEW_PAGE if page_break else WD_SECTION.CONTINUOUS)
    _set_section_orientation(doc, landscape=landscape)


def _pdf_content_inches(doc: Document) -> tuple[float, float]:
    """Область печати текущей секции."""
    if not doc.sections:
        return 6.0, 9.0
    return section_content_inches(doc.sections[-1])


def _fit_picture_inches(
    width_px: int,
    height_px: int,
    *,
    dpi: int,
    max_width_inches: float,
    max_height_inches: float,
) -> tuple[float, float]:
    w_in = width_px / dpi
    h_in = height_px / dpi
    if w_in <= 0 or h_in <= 0:
        return max_width_inches, max_height_inches
    scale = min(max_width_inches / w_in, max_height_inches / h_in, 1.0)
    return w_in * scale, h_in * scale


def _pdf_render_rect(page: fitz.Page) -> fitz.Rect:
    """Область рендера PDF: не обрезать контент из-за узкого CropBox."""
    rect = page.rect | page.mediabox
    try:
        bound = page.bound()
        if bound.width > 0 and bound.height > 0:
            rect |= bound
    except Exception:
        pass
    return rect


def _append_pdf_as_images(
    doc: Document,
    pdf_path: Path,
    *,
    dpi: int = 120,
) -> None:
    max_w_in, max_h_in = _pdf_content_inches(doc)
    matrix = fitz.Matrix(dpi / 72.0, dpi / 72.0)

    src = fitz.open(pdf_path)
    try:
        for i in range(len(src)):
            page = src[i]
            _ensure_page_orientation(doc, landscape=page.rect.width > page.rect.height)
            max_w_in, max_h_in = _pdf_content_inches(doc)
            pix = page.get_pixmap(matrix=matrix, clip=_pdf_render_rect(page), alpha=False)
            w_in, h_in = _fit_picture_inches(
                pix.width,
                pix.height,
                dpi=dpi,
                max_width_inches=max_w_in,
                max_height_inches=max_h_in,
            )
            bio = io.BytesIO(pix.tobytes("png"))
            bio.seek(0)
            par = doc.add_paragraph()
            run = par.add_run()
            run.add_picture(bio, width=Inches(w_in), height=Inches(h_in))
    finally:
        src.close()


def _append_raster_image(doc: Document, image_path: Path) -> None:
    image = DocxImage.from_file(str(image_path))
    horz_dpi = image.horz_dpi or 72
    vert_dpi = image.vert_dpi or horz_dpi
    w_in = image.px_width / horz_dpi
    h_in = image.px_height / vert_dpi
    _ensure_page_orientation(doc, landscape=w_in > h_in)
    max_w_in, max_h_in = _pdf_content_inches(doc)
    if w_in <= 0 or h_in <= 0:
        w_in, h_in = max_w_in, max_h_in
    scale = min(max_w_in / w_in, max_h_in / h_in, 1.0)

    par = doc.add_paragraph()
    run = par.add_run()
    run.add_picture(str(image_path), width=Inches(w_in * scale), height=Inches(h_in * scale))


def _append_vector_image(doc: Document, image_path: Path, *, dpi: int) -> None:
    matrix = fitz.Matrix(dpi / 72.0, dpi / 72.0)
    data = image_path.read_bytes()
    src = fitz.open(stream=data, filetype="svg")
    try:
        if len(src) == 0:
            raise ValueError("векторное изображение не содержит страниц")
        page = src[0]
        pix = page.get_pixmap(matrix=matrix, alpha=False)
        _ensure_page_orientation(doc, landscape=pix.width > pix.height)
        max_w_in, max_h_in = _pdf_content_inches(doc)
        w_in, h_in = _fit_picture_inches(
            pix.width,
            pix.height,
            dpi=dpi,
            max_width_inches=max_w_in,
            max_height_inches=max_h_in,
        )
        bio = io.BytesIO(pix.tobytes("png"))
        bio.seek(0)
        par = doc.add_paragraph()
        run = par.add_run()
        run.add_picture(bio, width=Inches(w_in), height=Inches(h_in))
    finally:
        src.close()


def _image_is_landscape(image_path: Path, *, dpi: int, warnings: list[str]) -> bool:
    ext = image_path.suffix.lower()
    try:
        if ext in RASTER_IMAGE_EXTS:
            image = DocxImage.from_file(str(image_path))
            horz_dpi = image.horz_dpi or 72
            vert_dpi = image.vert_dpi or horz_dpi
            return (image.px_width / horz_dpi) > (image.px_height / vert_dpi)
        if ext in VECTOR_IMAGE_EXTS:
            matrix = fitz.Matrix(dpi / 72.0, dpi / 72.0)
            src = fitz.open(stream=image_path.read_bytes(), filetype="svg")
            try:
                if len(src) == 0:
                    return False
                pix = src[0].get_pixmap(matrix=matrix, alpha=False)
                return pix.width > pix.height
            finally:
                src.close()
    except Exception as e:
        warnings.append(f"Не удалось определить ориентацию изображения ({image_path.name}): {e}")
    return False


def _append_image(doc: Document, image_path: Path, *, dpi: int) -> None:
    ext = image_path.suffix.lower()
    if ext in RASTER_IMAGE_EXTS:
        _append_raster_image(doc, image_path)
        return
    if ext in VECTOR_IMAGE_EXTS:
        _append_vector_image(doc, image_path, dpi=dpi)
        return
    raise ValueError(f"неподдерживаемый формат изображения: {image_path.suffix}")


def merge_to_docx_and_pdf(
    paths: list[Path],
    output_docx: Path,
    output_pdf: Path | None,
    *,
    page_break_between_parts: bool = True,
    insert_titles: bool = False,
    pdf_render_dpi: int = 120,
    libreoffice_executable: str | None = None,
) -> tuple[list[str], list[str]]:
    """
    Склеивает файлы по порядку `paths`.
    Возвращает (warnings, errors). errors непустой — часть шагов не выполнена.
    """
    warnings: list[str] = []
    errors: list[str] = []
    soffice = find_soffice(preferred=libreoffice_executable)

    work_items: list[tuple[Path, str]] = []
    image_items: list[tuple[Path, str]] = []
    temp_to_delete: list[Path] = []

    for p in paths:
        p = Path(p)
        if not p.is_file():
            warnings.append(f"Пропуск (файл не найден): {p}")
            continue
        ext = p.suffix.lower()
        if ext == ".doc":
            warnings.append(f"Пропуск .doc без чтения: {p.name}")
            continue
        if ext == ".docx":
            work_items.append((p, p.name))
            continue
        if ext == ".pdf":
            work_items.append((p, p.name))
            continue
        if ext in IMAGE_EXTS:
            image_items.append((p, p.name))
            continue
        warnings.append(f"Пропуск (неподдерживаемый тип): {p.name}")

    if not work_items and not image_items:
        errors.append("Нет ни одного поддерживаемого файла для склейки (.doc пропускаются).")
        return warnings, errors

    merged: Document | None = None
    composer: Composer | None = None
    part_sections: list[tuple[int, int, int]] = []
    global_page_offset = 0

    output_docx = Path(output_docx)
    try:
        for idx, (src_path, display_name) in enumerate(work_items):
            ext = src_path.suffix.lower()
            page_count = _source_page_count(src_path, soffice=soffice, warnings=warnings)

            if ext == ".pdf":
                if merged is None:
                    merged = Document()
                    composer = None
                section_start = 0 if not _has_body_content(merged) else len(merged.sections)
                with fitz.open(src_path) as src_pdf:
                    first_landscape = len(src_pdf) > 0 and src_pdf[0].rect.width > src_pdf[0].rect.height
                _start_new_part_section_for_orientation(
                    merged,
                    landscape=first_landscape,
                    page_break=page_break_between_parts,
                )
                if insert_titles:
                    merged.add_heading(display_name, level=2)
                _append_pdf_as_images(merged, src_path, dpi=pdf_render_dpi)
                part_sections.append((section_start, len(merged.sections) - 1, global_page_offset))
                global_page_offset += page_count
                continue

            if ext == ".docx":
                doc = Document(str(src_path))
                remove_empty_pages(doc)
                if insert_titles:
                    _prepend_heading(doc, display_name)
                if merged is not None:
                    section_start = len(merged.sections)
                    _start_new_part_section_like(merged, doc, page_break=page_break_between_parts)
                if merged is None:
                    merged = doc
                    composer = Composer(merged)
                    section_start = 0
                else:
                    if composer is None:
                        composer = Composer(merged)
                    composer.append(doc)
                part_sections.append((section_start, len(merged.sections) - 1, global_page_offset))
                global_page_offset += page_count
                continue

        for src_path, display_name in image_items:
            page_count = _source_page_count(src_path, soffice=soffice, warnings=warnings)
            if merged is None:
                merged = Document()
                composer = None
            section_start = 0 if not _has_body_content(merged) else len(merged.sections)
            landscape = _image_is_landscape(src_path, dpi=pdf_render_dpi, warnings=warnings)
            _start_new_part_section_for_orientation(
                merged,
                landscape=landscape,
                page_break=page_break_between_parts,
            )
            if insert_titles:
                merged.add_heading(display_name, level=2)
            try:
                _append_image(merged, src_path, dpi=pdf_render_dpi)
                part_sections.append((section_start, len(merged.sections) - 1, global_page_offset))
                global_page_offset += page_count
            except Exception as e:
                warnings.append(f"Изображение не вставлено ({src_path.name}): {e}")

        if merged is None:
            errors.append("Не удалось сформировать документ (неизвестная причина).")
            for t in temp_to_delete:
                t.unlink(missing_ok=True)
            return warnings, errors

        remove_empty_pages(merged)
        _apply_page_numbering(merged, part_sections)
        output_docx.parent.mkdir(parents=True, exist_ok=True)
        merged.save(str(output_docx))
    except Exception as e:
        errors.append(f"Ошибка при склейке или сохранении DOCX: {e}\n{traceback.format_exc()}")
        for t in temp_to_delete:
            t.unlink(missing_ok=True)
        return warnings, errors

    if output_pdf is not None:
        output_pdf = Path(output_pdf)
        output_pdf.parent.mkdir(parents=True, exist_ok=True)
        if not soffice:
            warnings.append(
                "LibreOffice (soffice) не найден — PDF не создан. "
                "Укажите полный путь к soffice.exe в настройках окна, "
                "переменную LIBREOFFICE_EXECUTABLE или установите LibreOffice."
            )
        else:
            try:
                docx_to_pdf(soffice, output_docx, output_pdf.parent)
                lo_out = output_pdf.parent / f"{output_docx.stem}.pdf"
                if lo_out.is_file():
                    if lo_out.resolve() != output_pdf.resolve():
                        shutil.move(str(lo_out), str(output_pdf))
                else:
                    errors.append(f"LibreOffice не создал PDF: {lo_out}")
            except Exception as e:
                errors.append(f"Экспорт PDF не удался: {e}")

    for t in temp_to_delete:
        try:
            t.unlink(missing_ok=True)
            if t.parent.is_dir() and not any(t.parent.iterdir()):
                t.parent.rmdir()
        except OSError:
            warnings.append(f"Не удалось удалить временный файл: {t}")

    return warnings, errors
