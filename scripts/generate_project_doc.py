# -*- coding: utf-8 -*-
"""Генерирует DOCX с описанием текущего состояния проекта GEO Documents."""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "GEO_Documents_project_v1.0.4.docx"


def main() -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    title = doc.add_heading("GEO Documents — описание проекта", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("Версия: 1.0.4 (коммит 9d8c913)")
    doc.add_paragraph("Дата документа: 3 июня 2026")
    doc.add_paragraph(
        "Назначение: десктопное приложение для склейки отчётных материалов "
        "(PDF, DOCX, изображения) в единый DOCX с опциональным экспортом в PDF."
    )

    sections: list[tuple[str, list[str]]] = [
        (
            "1. Технологии",
            [
                "Python 3.12",
                "PyQt6 — графический интерфейс",
                "python-docx — работа с DOCX",
                "docxcompose — объединение DOCX-фрагментов",
                "PyMuPDF (fitz) — растеризация PDF и SVG",
                "LibreOffice (soffice.exe) — экспорт DOCX → PDF",
                "PyInstaller — сборка GEO_Documents.exe для Windows",
            ],
        ),
        (
            "2. Поддерживаемые форматы",
            [
                "DOCX — вставляется с сохранением ориентации и размеров секции исходника",
                "PDF — каждая страница вставляется как изображение (настраиваемый DPI)",
                "PNG, JPG, JPEG — растровые изображения",
                "SVG, DVG — векторные изображения (растеризация через PyMuPDF)",
                "DOC — не читается, файл пропускается с предупреждением",
                "DWG и прочие форматы — не поддерживаются",
            ],
        ),
        (
            "3. Интерфейс приложения",
            [
                "Выбор папки и сканирование файлов (.pdf, .docx, изображения)",
                "Список с drag-and-drop, кнопки «вверх/вниз», добавление и удаление файлов",
                "Автосортировка по правилам геоотчётов (разделы 1.x, приложения А., графика 3.N)",
                "Параметры: разрыв страницы между частями, заголовки с именем файла, DPI для PDF",
                "Путь к soffice.exe (сохраняется в QSettings)",
                "Результат: имя.docx и имя.pdf в выбранной папке",
            ],
        ),
        (
            "4. Логика склейки (merger.py)",
            [
                "Порядок файлов — как в списке пользователя (после автосортировки при сканировании папки)",
                "PDF и DOCX обрабатываются в порядке списка",
                "Изображения (png/jpg/svg/…) собираются отдельно и вставляются в конец итогового документа",
                "Между фрагментами — разрыв страницы (если включён)",
                "DOCX: при другой геометрии секции создаётся новая секция с копированием размеров/полей исходника",
                "PDF/картинки: альбомная ориентация секции, если ширина больше высоты",
                "Экспорт PDF через LibreOffice, если найден soffice.exe",
            ],
        ),
        (
            "5. Модуль docx_fit.py",
            [
                "Содержит функции подгонки таблиц, рисунков, подписей «Составил:», удаления пустых страниц",
                "В версии 1.0.4 не вызывается из merger.py при склейке — модуль подготовлен, но не подключён к основному потоку",
            ],
        ),
        (
            "6. Автосортировка (file_sorter.py)",
            [
                "Приоритет 0: числовые разделы (1.2, 1.4.1)",
                "Приоритет 1: приложения буквой (А., Г.3, В.Ведомость)",
                "Приоритет 2: графика 3.N и псевдонимы э./ю./я.",
                "Остальные файлы — в конце списка",
            ],
        ),
        (
            "7. Сборка и запуск",
            [
                "Запуск из исходников: python -m geo_documents",
                "Сборка: build_windows.bat или PyInstaller GEO_Documents.spec",
                "Результат сборки: dist\\GEO_Documents.exe",
                "Зависимости: requirements.txt, для сборки — requirements-build.txt",
            ],
        ),
        (
            "8. Структура проекта",
            [
                "launcher.py — точка входа для exe",
                "geo_documents/main.py — QApplication",
                "geo_documents/window.py — главное окно",
                "geo_documents/merger.py — склейка",
                "geo_documents/docx_fit.py — обработка DOCX (не в merge)",
                "geo_documents/file_sorter.py — сортировка имён",
                "geo_documents/libreoffice.py — поиск soffice и конвертация",
            ],
        ),
        (
            "9. Ограничения",
            [
                "Файлы .doc не конвертируются автоматически в текущей версии UI (только предупреждение)",
                "Без LibreOffice итоговый PDF не создаётся",
                "Сложные DOCX с плавающими объектами могут требовать ручной проверки в Word",
                "Точная вёрстка зависит от исходных секций Word; при склейке разных форматов листа создаются отдельные секции",
            ],
        ),
    ]

    for heading, bullets in sections:
        doc.add_heading(heading, level=1)
        for item in bullets:
            doc.add_paragraph(item, style="List Bullet")

    doc.save(str(OUT))
    print(OUT)


if __name__ == "__main__":
    main()
